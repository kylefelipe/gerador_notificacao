# -*- coding: utf-8 -*-

##################################################################################
#     Script python para gerar vários documentos usando um DOCX como template
# e uma planilha XLSX como base de dados.
# é necessário ter um template preparado antes.
#
# Author: Kyle Felipe Vieira Roberto
# Data de Criação: quarta-feira 19 de julho de 2017
#
##################################################################################

from openpyxl import load_workbook
import funcoes_base_dados
import funcoes_template
from docx import Document
from docx import table
import re

# Código de auxílio
remover = "[-,./\()]"

# Trabalhando a base de dados XLSX

while True:
    try:
        solicita_caderno = str(input("Digite o nome do arquivo com os Dados:  "))
        caderno = load_workbook(filename=solicita_caderno + '.xlsx')
        break
    except:
        print('Arquivo não encontrado, confira o nome do arquivo com a base de dados.')
folha = caderno[funcoes_base_dados.escolhe_folha(caderno)]
cnpjcpf = funcoes_base_dados.listar_documentos_contribuintes(folha)
lista_documentos = funcoes_base_dados.listar_cnpjcpf_contribuintes(folha)
total = len(cnpjcpf)  # Quantidade de notificaçeõs a serem feitas

# Solicita o template para as notificações
while True:
    try:
        arquivo_base = input("Digite o nome do arquivo DOCX padrão: ")
        documento = Document(arquivo_base + '.docx')
        break
    except:
        print('Arquivo não encontrado, confira o nome do arquivo do template.')

# Solicita dados para a notificação

print('Total de notificações: {}' .format(total))
data_notificacao = input("Digite a data de emissão com o mês por extenso: ")
while True:
    try:
        notificacao_inicial = int(input("Qual o número da primeira notificação: "))  # Pega o número da primeira notificação
        break
    except:
        print('O número da notificação precisa ser um número inteiro')

# Criando as notificações

for i in range(total):
    num_notificacao = notificacao_inicial + i  # Criando o número da notificação
    documento = Document(arquivo_base+'.docx')  # Abrindo o arquivo padrão
    cnpj = cnpjcpf[i]
    nome_limpo = re.sub(remover, '', lista_documentos[cnpj])  # Limpa o nome da empresa.
    itens_notificacao = funcoes_base_dados.dados_empresa_tabela(cnpj, folha)

    # Listando os paragrafos do arquivo template
    paragrafo = list(documento.paragraphs)

    # Modifica o texto do paragrafo 0 - número da notificação
    numnot_para = paragrafo[0].text
    paragrafo[0].text = numnot_para.format(num_notificacao)

    # Modifica o texto do paragrafo 3 - data de emissão
    data_emissao = paragrafo[3].text
    paragrafo[3].text = data_emissao.format(data_notificacao)

    # Modifica o texto do paragrafo 5 - Nome da empresa
    nome_empresa = paragrafo[5].text
    paragrafo[5].text = nome_empresa.format(nome_limpo)


    # Modifica o texto do paragrafo 6 - CNPJ
    cnpj_empresa = paragrafo[6].text
    paragrafo[6].text = cnpj_empresa.format(cnpj)

    # Trabalhando a tabela da Notificação

    tabelas_dados = list(documento.tables)
    tabela_d = funcoes_template.popula_tabela(tabelas_dados, itens_notificacao, documento)
    tabelas_dados[0]=tabela_d
    print('Gerando notificação nº {0} - {1}' .format(num_notificacao, nome_limpo))
    documento.save('NOTIFICAÇÃO nº {0} - {1} .docx' .format(num_notificacao, nome_limpo))


